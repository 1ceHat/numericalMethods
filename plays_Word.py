from math import *
from docx import *
from docx.enum.style import *
from docx.enum.text import *
from docx.enum import *
from docx.shared import *
doc = Document()
style1 = doc.styles.add_style('subscript', WD_STYLE_TYPE.CHARACTER)
style1.font.subscript = True
style2 = doc.styles.add_style('supscript', WD_STYLE_TYPE.CHARACTER)
style2.font.superscript = True
doc.add_paragraph()

def user_y(f, x, y, num):
    res = round(eval(f), num)
    return res

# Метод преобразует символы идущие после _( и ^( в индекс и степень соответственно
def make_subscript_supscript(prg):
    prg = prg
    words = prg.runs[-1].text
    prg.runs[-1].clear()

    while words.count('_')!=0 or words.count('^')!=0:
        ind_subsc = words.find('_')
        ind_supsc = words.find('^')

        if (ind_subsc<ind_supsc or ind_supsc == -1) and ind_subsc!=-1:
            run = prg.add_run(words[:ind_subsc])
            words = words[ind_subsc:]
            ind_Lprth = words.find('(')
            ind_Rprth = words.find(')')
            run = prg.add_run(words[ind_Lprth+1:ind_Rprth], style='subscript')
            words = words[ind_Rprth+1:]
        elif (ind_supsc<ind_subsc or ind_subsc == -1) and ind_supsc!=-1:
            run = prg.add_run(words[:ind_supsc])
            words = words[ind_supsc+1:]
            ind_Lprth = words.find('(')
            ind_Rprth = words.find(')')
            run = prg.add_run(words[ind_Lprth+1:ind_Rprth], style='supscript')
            words = words[ind_Rprth+1:]
    run = prg.add_run(words)

# Метод вычисляет таблицу значений с точностью num по методу Эйлера
def calc_Eyil(f, y, h, num):
    n = int(1 / h + 1)
    for i in range(2):
        for j in range(1, n):
            if i == 0:
                y[0].append(round(h * j, 2))

            elif i == 1:
                fi1 = user_y(f, y[0][j - 1], y[1][j - 1], num)
                fi = 1 + 15 * y[1][j - 1] * cos(y[0][j - 1]) - y[1][j - 1] ** 2  # <---- ТУТ МЕНЯЕШЬ ФУНКЦИЮ

                y[1].append(round(y[1][j - 1] + h * fi1, num))

# Метод выводит вычисления таблицы значений по методу Эйлера
def print_calcE(f, y, h, doc):
    n = int(1 / h + 1)

    prg = doc.paragraphs[-1]
    run = prg.add_run('Формула, используемая в этом методе:\n'
                      '\ty_(n+1) = y_(n) + h*F(x_(n) , y_(n)),\n'
                      'где\th - размер шага,\n'
                      '\tF(x_(n), y_(n)) - значение производной.\n\n'
                      'Вычислим y_(i) :\n'
                      'y_(0) = ' + str(int(y[1][0])))

    make_subscript_supscript(prg)

    for i in range(1, n):
        # Пропишем предварительно формулу
        f1 = f.replace('y', 'y_(' + str(i - 1)+')')
        f1 = f1.replace('x', 'x_(' + str(i - 1)+')')
        res1 = 'y_(' + str(i - 1) + ') + h * (' + f1 + ') = '

        # Подставим в формулу значения
        f2 = f.replace('y', str(y[1][i - 1]))
        f2 = f2.replace('x', str(y[0][i - 1]))
        res2 = str(y[1][i - 1]) + ' + ' + str(h) + ' * (' + f2 + ') ≈ '

        # Выводим итоговое выражение
        res = 'y_(' + str(i) + ') = ' + res1 + res2 + str(y[1][i])

        prg = doc.add_paragraph()
        run = prg.add_run(res)
        make_subscript_supscript(prg)
    prg = doc.add_paragraph()

# Тут таблица значений для Эйлера
def print_Eyil(y, step, doc):
    tab = doc.add_table(2,len(y[0])+1, 'Table Grid')
    prg = tab.cell(0,0).paragraphs[-1]
    prg.add_run('x')
    prg.add_run('i', style = 'subscript')
    prg = tab.cell(1,0).paragraphs[-1]
    prg.add_run('y')
    prg.add_run('i', style = 'subscript')

    for i in range(2):
        for j in range(len(y[0])):
            tab.cell(i, j+1).paragraphs[-1].add_run(str(y[i][j]))
            tab.rows[i].height = Cm(1)
            tab.columns.width = Cm(2)

    print('x_i: ', *y[0])
    print('y_i: ', *y[1], '\n')
    '''if step == 0.1:
        DraGraphs(y,step,"Эйлера")'''

# Метод считает таблицу значений по методу Эйлера-Коши
def calc_EyilKosh(f, y, h, num):
    n = int(1/h+1)

    for i in range(2):
        for j in range(1,n):
            if i==0:
                y[0].append(round(h*j, 2))

            elif i==1:
                '''=======Заполняем yi======='''
                # Считаем функцию F(xi-1, yi-1)
                fi1 = user_y(f, y[0][j-1], y[2][j-1], num)
                # y[0][] - это х
                # y[1][] - это y`
                # y[2][] - это y
                #F1.append(fi1)

                '''== Cчитаем функцию F(xi, y`i) =='''
                # Вычисляем y`
                resY1 = y[2][j-1] + h*fi1
                y[1].append(round(resY1, num))

                # Вычисляем функцию F(xi, y`i)
                fi2 = user_y(f, y[0][j], y[1][j], num)
                # y[0][] - это х
                # y[1][] - это y`
                # y[2][] - это y
                #F2.append(fi2)

                # Посчитаем сумму функций F(xi-1, yi-1) и F(xi, y`i)
                sumF = fi1 + fi2
                resY2 = y[2][j-1] + h * 0.5 * sumF
                y[2].append(round(resY2, num))

# Метод выводит вычисления таблицы значений по методу Эйлера-Коши
def print_calcEK(f, y, h, num, doc):
    n = int(1 / h + 1)

    prg = doc.paragraphs[-1]
    run = prg.add_run('Формулы, используемые в этом методе:\n'
          '\ty`_(i) = y_(i-1) + h * F(x_(i-1), y_(i-1)),\n'
          'где\ty`_(i) - первое значение производной,\n'
          '\th - размер шага,\n'
            '\tF(x_(i-1), y_(i-1)) - предыдущее значение производной;\n\n'
          '\ty_(i) = y_(i-1) + h * 0.5 * (F(x_(i), y`_(i)) + F(x_(i-1), y_(i-1)))\n'
          'где\ty_(i) - значение производной после пересчёта,\n'
          '\tF(x_(i), y`_(i)) - значение производной от первого значения производной,\n'
          '\tF(x_(i-1), y_(i-1)) - предыдущее значение производной.\n\n'
          'Вычислим y_(i):\n'
          'y_(0) =' + str(int(y[1][0])))
    make_subscript_supscript(prg)

    ysi = 'y_(i-1) + h * F(x_(i-1), y_(i-1))'
    yi = 'y_(i-1) + h * 0.5 * (F(x_(i), y`_(i)) + F(x_(i-1), y_(i-1)))'
    for i in range(1,n):

        '''=======Расчёт y`i======='''
        # Пропишем предварительно формулу
        f1 = ysi.replace('y_(i-1)', 'y_('+str(i-1)+')')
        f1 = f1.replace('x_(i-1)', 'x_('+str(i-1)+')')
        res1 = f1 + ' = '

        # Подставим в формулу значения
        f2 = f1.replace('y_('+str(i-1)+')', str(y[2][i-1]))
        f2 = f2.replace('x_('+str(i-1)+')', str(y[0][i-1]))
        f2 = f2.replace('h', str(h))
        res2 = f2 + ' ≈ '

        # Выводим итоговое выражение
        res = 'y`_(' + str(i) + ')\t= ' + res1 + res2 + str(y[1][i])
        prg = doc.add_paragraph()
        prg.add_run(res)
        prg.paragraph_format.tab_stops.add_tab_stop(Cm(0.5))
        make_subscript_supscript(prg)



        '''=======Расчёт yi======='''
        # Пропишем предварительно формулу
        f1 = yi.replace('y_(i-1)', 'y_(' + str(i - 1)+')')
        f1 = f1.replace('x_(i-1)', 'x_(' + str(i - 1)+')')
        f1 = f1.replace('x_(i), y`_(i)', 'x_('+str(i)+'), y`_('+str(i)+')')
        res1 = f1 + ' = '

        # Подставим в формулу значения
        f2 = f1.replace('y_(' + str(i - 1)+')', str(y[2][i - 1]))
        f2 = f2.replace('x_(' + str(i - 1)+')', str(y[0][i - 1]))
        f2 = f2.replace('y`_(' + str(i)+')', str(y[1][i]))
        f2 = f2.replace('x_(' + str(i)+')', str(y[0][i]))
        f2 = f2.replace('h', str(h))
        res2 = f2 + ' =\n\t= '

        # Небольшое промежуточное значение, вставим значения функций
        fi1 = user_y(f, y[0][i], y[1][i], num)
        fi2 = user_y(f, y[0][i-1], y[2][i-1], num)
        f3 = f2.replace('F('+str(y[0][i])+', '+str(y[1][i])+')', str(round(fi1, num)))
        f3 = f3.replace('F('+str(y[0][i-1])+', '+str(y[2][i-1])+')', str(round(fi2, num)))
        res3 = f3 + ' ≈ '

        # Выводим итоговое выражение
        res = 'y_(' + str(i) + ')\t= ' + res1 + res2 + res3 + str(y[2][i])
        prg = doc.add_paragraph()
        prg.add_run(res)
        prg.paragraph_format.tab_stops.add_tab_stop(Cm(0.5))
        make_subscript_supscript(prg)
    prg = doc.add_paragraph()

# Метод выводит таблицу значений по методу Эйлера-Коши
def print_EyilKosh(y, step, doc):
    tab = doc.add_table(3, len(y[0]) + 1, 'Table Grid')
    prg = tab.cell(0, 0).paragraphs[-1]
    prg.add_run('x')
    prg.add_run('i', style='subscript')
    prg = tab.cell(1, 0).paragraphs[-1]
    prg.add_run('y')
    prg.add_run('i', style='subscript')
    prg.add_run('\'', style='supscript')
    prg = tab.cell(2, 0).paragraphs[-1]
    prg.add_run('y')
    prg.add_run('i', style='subscript')

    for i in range(3):
        for j in range(len(y[0])):
            tab.cell(i, j+1).paragraphs[-1].add_run(str(y[i][j]))
            tab.rows[i].height = Cm(1)
            tab.columns.width = Cm(2)

    print('xi: ', *y[0])
    print('yi\': ', *y[1])
    print('yi: ', *y[2], '\n')
    '''if step == 0.1:
        DraGraphs(y,step,"Эйлера-Коши")'''

# Метод считает таблицу значений по методу Рунге-кутты
def calc_RungeKutta(f, y, h, num):
    n = int(1/h+1)
    for i in range(2):
        for j in range(1, n):
            if i==0:
                y[0].append(round(h * j, 2))
            elif i==1:
                k1 = user_y(f, y[0][j - 1], y[1][j - 1], num)
                k2 = user_y(f, y[0][j - 1] + h/2, y[1][j - 1] + k1*h/2, num)
                k3 = user_y(f, y[0][j - 1] + h/2, y[1][j - 1] + k2*h/2, num)
                k4 = user_y(f, y[0][j - 1] + h, y[1][j - 1] + k3*h, num)

                yi = y[1][j-1] + h/6*(k1 + 2*k2 + 2*k3 + k4)
                y[1].append(round(yi, num))

# Выводит вычисления таблицы по методу Рунге-Кутты
def print_calcRK(fc, y, h, num, doc):
    n = int(1 / h + 1)
    k_1 = 'F(x_(i), y_(i))'
    k_2 = 'F(x_(i) + h/2, y_(i) + h/2*k_(i)^(1))'
    k_3 = 'F(x_(i) + h/2, y_(i) + h/2*k_i^(2))'
    k_4 = 'F(x_(i) + h, y_(i) + h*k_(i)^(3))'
    yi = 'y_(i) + h/6*(k_(i)^(1) +  2*k_(i)^(2) + 2*k_(i)^(3) + k_(i)^(4))'
    f = fc.replace('**','^')

    prg = doc.paragraphs[-1]
    run = prg.add_run('Формулы, используемые в этом методе:\n'
          '\ty_(i+1) = y_(i) + h/6*(k_(i)^(1) + 2*k_(i)^(2) + 2*k_(i)^(3) + k_(i)^(4)),\n'
          'где\ty_(i+1) - искомое значение производной в точке,\n'
          '\tk_(i) - коэффициенты;\n\n'
          '\tk_(i)^(1) = F(x_(i), y_(i)),\n'
          '\tk_(i)^(2) = F(x_(i) + h/2, y_(i) + h/2*k_(i)^(1)),\n'
          '\tk_(i)^(3) = F(x_(i) + h/2, y_(i) + h/2*k_(i)^(2)),\n'
          '\tk_(i)^(4) = F(x_(i) + h, y_(i) + h*k_(i)^(3)).\n\n'
          'Вычислим y_(i):\n'
          'y_(0) =' + str(int(y[1][0])) + '\n')
    make_subscript_supscript(prg)

    for i in range(1, n):
        # Найдём коэффициенты

        # k1
        # Напишем формулу
        f1 = k_1.replace('x_(i), y_(i)', 'x_('+str(i-1)+') ,  y_('+str(i-1)+')')
        # Подставим значения в формулу
        fi = f.replace('x', str(y[0][i-1]))
        fi = fi.replace('y', str(y[1][i-1]))
        # Считаем значение k
        k1 = user_y(fc, y[0][i - 1], y[1][i - 1], num)
        res1 = 'k_(1) = ' + f1 + ' = ' + fi + ' ≈ ' + str(k1)
        prg = doc.add_paragraph()
        run = prg.add_run(res1)
        make_subscript_supscript(prg)

        # k2
        # Напишем формулу
        f2 = k_2.replace('x_(i), ', 'x_(' + str(i-1)+')')
        f2 = f2.replace('y_(i)', 'y_(' + str(i-1)+')')
        f2 = f2.replace('h/2', str(h/2))
        f2 = f2.replace('k_(i)', 'k_('+str(i-1)+')')
        # Подставим значения в формулу
        fi = f.replace('x', str(round(y[0][i-1]+h/2, 4)))
        fi = fi.replace('y', str(round(y[1][i-1]+h/2*k1, 4)))
        k2 = user_y(fc, y[0][i - 1] + h / 2, y[1][i - 1] + k1 * h / 2, num)
        res2 = 'k_(2) = ' + f2 + ' = ' + fi + ' ≈ ' + str(k2)
        prg = doc.add_paragraph()
        run = prg.add_run(res2)
        make_subscript_supscript(prg)

        # k3
        # Напишем формулу
        f3 = k_3.replace('x_(i), ', 'x_(' + str(i - 1) + ')')
        f3 = f3.replace('y_(i)', 'y_(' + str(i - 1) + ')')
        f3 = f3.replace('h/2', str(h/2))
        f3 = f3.replace('k_(i)', 'k_(' + str(i - 1) + ')')
        # Подставим значения в формулу
        fi = f.replace('x', str(round(y[0][i - 1] + h / 2, 4)))
        fi = fi.replace('y', str(round(y[1][i - 1] + h / 2 * k2, 4)))
        k3 = round(user_y(fc, y[0][i - 1] + h / 2, y[1][i - 1] + k2 * h / 2, num), num)
        res3 ='k_(3) = ' + f3 + ' = ' + fi + ' ≈ ' + str(k3)
        prg = doc.add_paragraph()
        run = prg.add_run(res3)
        make_subscript_supscript(prg)

        # k4
        # Напишем формулу
        f4 = k_4.replace('x_(i), ', 'x_(' + str(i - 1)+')')
        f4 = f4.replace('y_(i)', 'y_(' + str(i - 1)+')')
        f4 = f4.replace('h', str(h))
        f4 = f4.replace('k_(i)', 'k_(' + str(i - 1)+')')
        # Подставим значения в формулу
        fi = f.replace('x', str(round(y[0][i - 1] + h, 4)))
        fi = fi.replace('y', str(round(y[1][i - 1] + h * k3, 4)))
        k4 = round(user_y(fc, y[0][i - 1] + h, y[1][i - 1] + k3 * h, num),4)
        res4 = 'k_(4) = ' + f4 + ' = ' + fi + ' ≈ ' + str(k4)
        prg = doc.add_paragraph()
        run = prg.add_run(res4)
        make_subscript_supscript(prg)

        # yi
        # Напишем формулу
        yi1 = yi.replace('y_(i)', 'y_' + str(i-1)+')')
        yi1 = yi1.replace('k_(i)', 'k_' + str(i-1)+')')
        resy1 = yi1 + ' = '
        # Подставим значений
        yi2 = yi.replace('y_(i)', str(y[1][i-1]))
        yi2 = yi2.replace('k_(' + str(i - 1) + '^(1)', str(k_1))
        yi2 = yi2.replace('k_(' + str(i - 1) + '^(2)', str(k_2))
        yi2 = yi2.replace('k_(' + str(i - 1) + '^(3)', str(k_3))
        yi2 = yi2.replace('k_(' + str(i - 1) + '^(4)', str(k_4))
        yi2 = yi2.replace('h', str(h))
        resy2 = yi2 + ' ≈ '
        # Выводим итоговый результат:
        res = 'y_(' + str(i) + ') = ' + resy1 + resy2 + str(y[1][i]) + '\n'
        prg = doc.add_paragraph()
        run = prg.add_run(res)
        make_subscript_supscript(prg)
    prg = doc.add_paragraph()

# Выводит таблицу значений по методу Рунге-Кутты
def print_RungeKutta(y,step):
    tab = doc.add_table(2, len(y[0]) + 1, 'Table Grid')
    prg = tab.cell(0, 0).paragraphs[-1]
    prg.add_run('x')
    prg.add_run('i', style='subscript')
    prg = tab.cell(1, 0).paragraphs[-1]
    prg.add_run('y')
    prg.add_run('i', style='subscript')

    for i in range(2):
        for j in range(len(y[0])):
            ran = tab.cell(i, j+1).paragraphs[-1].add_run(str(y[i][j]))
            ran.font.size = Pt(10)
            tab.rows[i].height = Cm(1)
            tab.columns.width = Cm(2)
    '''if step == 0.1:
        DraGraphs(y,step,"Рунге-Кутты")'''

'''f = input('===Обозначения==='
          '\nсинус - sin(x)'
          '\nкосинус - cos(x)'
          '\nкорень - ^(1/n)'
          '\nстепень - ^n'
          '\n=================\n'
          'Введи функцию y`:\n').replace(',', '.')'''
f = '1 + 0.2 * y * sin(x) - y^(2)'
'''x0 = float(input('Введи x0: ').replace(',', '.'))
y0 = float(input('Введи у0: ').replace(',', '.'))'''
x0 = 0
y0 = 0

'''h1 = float(input('Введи первый шаг (h): ').replace(',', '.'))
h2 = float(input('Введи второй шаг (h): ').replace(',', '.'))'''
h1 = 0.2
h2 = 0.1

minh = min(h1,h2)
h1 = max(h1, h2)
h2 = minh

# Навели красоту в функции

f = f.replace(' ', '')
f = f.replace('*', ' * ')
f = f.replace('/', ' / ')
f = f.replace('+', ' + ')
f = f.replace('-', ' - ')
fc = f.replace('^', '**') # Такое представление функции надо заносить в calc_

y1 = [[x0],[y0]]
y2 = [[x0],[y0]]
print('Метод Эйлера')
#numEyil = int(input('До скольки знаков после запятой округлять? ').replace(',', '.'))
numEyil = 3
# Считаем для первого шага

calc_Eyil(fc, y1, h1, numEyil)
print_calcE(f, y1, h1, doc)
print_Eyil(y1, h1, doc)
doc.add_page_break()

#Считаем для второго шага
calc_Eyil(fc, y2, h2, numEyil)
print_calcE(f, y2, h2, doc)
print_Eyil(y2, h2, doc)
doc.add_page_break()

'''
    =========Метод Эйлера-Коши=========
    =======Считаем значения=======
'''
y3 = [[x0],[0],[y0]]
y4 = [[x0],[0],[y0]]
print('\nМетод Эйлера-Коши')
#numEyilKosh = int(input('До скольки знаков после запятой округлять? ').replace(',', '.'))
numEyilKosh = 3
# Считаем для первого шага
calc_EyilKosh(fc, y3, h1, numEyilKosh)
print_calcEK(fc, y3, h1, numEyilKosh, doc)
print_EyilKosh(y3, h1, doc)
doc.add_page_break()

#Считаем для второго шага
calc_EyilKosh(fc, y4, h2, numEyilKosh)
print_calcEK(fc, y4, h2, numEyilKosh, doc)
print_EyilKosh(y4, h2, doc)

y5 = [[x0],[y0]]
y6 = [[x0],[y0]]
print('\nМетод Рунге-Кутты')
#numRunge = int(input('До скольки знаков после запятой округлять? ').replace(',', '.'))
numRunge = 4
# Считаем для первого шага
calc_RungeKutta(fc, y5, h1, numRunge)
print_calcRK(fc, y5, h1, numRunge, doc)
print_RungeKutta(y5, h1)

#Считаем для второго шага
calc_RungeKutta(fc, y6, h2, 4)
print_calcRK(fc, y6, h2, 4, doc)
print_RungeKutta(y6, h2)

doc.save('tests.docx')